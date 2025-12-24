# =============================================
# BEAST MODE LINKEDIN SCRAPER (Class-Based)
# Transaction-safe Excel + Word + PDF + DB
# =============================================

import asyncio
import httpx
from bs4 import BeautifulSoup
import re
import random
import os
from datetime import datetime, timedelta
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.colors import blue
import tempfile
import shutil
import psycopg2

# =============================================
# CONFIG
# =============================================
KEYWORDS = ["DevOps Engineer", "Java Backend Developer"]
LOCATION = "United States"
MAX_RETRY = 3
SOURCE_PORTAL = "LinkedIn"


# =============================================
# DATABASE CLIENT (Transaction-Safe)
# =============================================
class DBClient:
    def __init__(self):
        self.conn = psycopg2.connect(
            host="localhost",
            port=5432,
            dbname="jobs_scraper",
            user="postgres",
            password="admin"
        )
        self.cur = self.conn.cursor()

    # -----------------------------
    # Transaction Control
    # -----------------------------
    def start_transaction(self):
        self.conn.autocommit = False

    def commit_transaction(self):
        self.conn.commit()
        self.conn.autocommit = True

    def rollback_transaction(self):
        self.conn.rollback()
        self.conn.autocommit = True

    # -----------------------------
    # Run Log
    # -----------------------------
    def log_run_start(self, keyword, source_portal):
        sql = """
        INSERT INTO scraper_run_log (keyword, source_portal, run_date, start_time, create_id)
        VALUES (%s, %s, CURRENT_DATE, NOW(), 'SCRAPER')
        RETURNING run_id;
        """
        self.cur.execute(sql, (keyword, source_portal))
        run_id = self.cur.fetchone()[0]
        self.conn.commit()
        print(f"🟢 Run started for keyword '{keyword}' | Run ID: {run_id}")
        return run_id

    def log_run_end(self, run_id, total):
        sql = """
        UPDATE scraper_run_log
        SET end_time = NOW(),
            total_jobs_scraped = %s,
            update_date = NOW(),
            update_id = 'SCRAPER'
        WHERE run_id = %s;
        """
        self.cur.execute(sql, (total, run_id))
        self.conn.commit()
        print(f"🟢 Run ended | Run ID: {run_id} | Total jobs: {total}")

    # -----------------------------
    # Upsert job_master
    # -----------------------------
    def upsert_master(self, job, source_portal):
        sql = """
            INSERT INTO job_master
            (job_title, company_name, location, posted_date, keyword, job_url, mobile_url,
             source_portal, create_id, create_date)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'SCRAPER', NOW())
            RETURNING job_id;
        """
        values = (
            job["Title"],
            job["Company"],
            job["Location"],
            job["Date Posted"],
            job["Keyword"],
            job["Job Link"],
            job["Mobile Link"],
            source_portal
        )
        self.cur.execute(sql, values)
        return self.cur.fetchone()[0]

    # -----------------------------
    # Archive master to history
    # -----------------------------
    def archive_master_to_history(self):
        sql = """
        INSERT INTO job_daily_history
        (job_id, snapshot_date, keyword, job_title, company_name, location,
         posted_date, job_url, mobile_url, source_portal, create_id)
        SELECT job_id, CURRENT_DATE, keyword, job_title, company_name, location,
               posted_date, job_url, mobile_url, source_portal, 'SCRAPER'
        FROM job_master;
        """
        self.cur.execute(sql)
        self.conn.commit()

    # -----------------------------
    # Clear today's job_master
    # -----------------------------
    def clear_master(self):
        self.cur.execute("""
            DELETE FROM job_master 
            WHERE create_date < CURRENT_DATE;
        """)
        self.conn.commit()
        print("🧹 Cleared job_master for today's run")

    # -----------------------------
    # Close connection
    # -----------------------------
    def close(self):
        self.cur.close()
        self.conn.close()
        print("🟢 Database connection closed")


# =============================================
# LINKEDIN SCRAPER
# =============================================
class LinkedInScraper:
    def __init__(self, location=LOCATION):
        self.location = location

    @staticmethod
    def posted_within_last_week(date_str):
        try:
            post_date = datetime.fromisoformat(date_str.replace("Z", ""))
            return post_date >= datetime.now() - timedelta(days=7)
        except:
            return False

    @staticmethod
    def get_mobile_link(job_link):
        match = re.search(r"/jobs/view/.*?-(\d+)", job_link)
        return f"https://www.linkedin.com/jobs/view/{match.group(1)}" if match else job_link

    @staticmethod
    def load_previous_ids(keyword):
        yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y%m%d")
        filename = f"{keyword.replace(' ', '')}/{keyword.replace(' ', '')}_{yesterday}.xlsx"
        if not os.path.exists(filename):
            return set()
        previous_ids = set()
        try:
            wb = load_workbook(filename)
            ws = wb.active
            for row in ws.iter_rows(min_row=2, values_only=True):
                job_link = row[5]  # Job Link column
                m = re.search(r"/jobs/view/.*?-(\d+)", job_link)
                if m:
                    previous_ids.add(m.group(1))
        except:
            pass
        return previous_ids

    async def fetch_jobs_for_keyword(self, client, keyword):
        url = "https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                          "(KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36"
        }

        job_postings = []
        seen_ids_today = set()
        seen_ids_yesterday = self.load_previous_ids(keyword)

        print(f"\n🚀 Starting scrape for keyword: {keyword}")
        print(f"📌 Loaded {len(seen_ids_yesterday)} previous job IDs (for dedupe)")

        page = 0
        retry = 0

        while True:
            params = {
                "keywords": keyword,
                "location": self.location,
                "sortBy": "R",
                "f_TPR": "r86400",
                "start": page * 25
            }

            print(f"🔄 Fetching page {page}...")

            try:
                resp = await client.get(url, headers=headers, params=params)

                if resp.status_code == 429:
                    wait_time = random.randint(45, 90)
                    print(f"🚫 Rate limited (429). Sleeping {wait_time} sec...")
                    await asyncio.sleep(wait_time)
                    continue

                if resp.status_code != 200:
                    print(f"❌ HTTP {resp.status_code} on page {page}")
                    retry += 1
                    if retry > MAX_RETRY:
                        print("❌ Too many failures. Stopping scraper.")
                        break
                    await asyncio.sleep(5)
                    continue

                retry = 0
                soup = BeautifulSoup(resp.content, "lxml")
                job_cards = soup.select("li")

                if not job_cards:
                    print(f"ℹ️ No more jobs found for: {keyword}")
                    break

                page += 1

                for job in job_cards:
                    try:
                        title_tag = job.find("h3")
                        company_tag = job.find("h4")
                        location_tag = job.find("span", class_="job-search-card__location")
                        time_tag = job.find("time")
                        link_tag = job.find("a", href=True)

                        if not all([title_tag, company_tag, location_tag, time_tag, link_tag]):
                            continue

                        job_link = link_tag["href"]
                        mobile_link = self.get_mobile_link(job_link)

                        job_id_match = re.search(r"/jobs/view/.*?-(\d+)", job_link)
                        if not job_id_match:
                            continue
                        job_id = job_id_match.group(1)

                        if job_id in seen_ids_today or job_id in seen_ids_yesterday:
                            continue

                        seen_ids_today.add(job_id)

                        title = title_tag.text.strip()
                        company = company_tag.text.strip()
                        location = location_tag.text.strip()
                        date_posted = time_tag.get("datetime", "")

                        if not self.posted_within_last_week(date_posted):
                            continue

                        job_postings.append({
                            "Title": title,
                            "Company": company,
                            "Location": location,
                            "Date Posted": date_posted,
                            "Keyword": keyword,
                            "Job Link": job_link,
                            "Mobile Link": mobile_link
                        })

                    except Exception as e:
                        print(f"⚠️ Parse Error: {e}")
                        continue

                await asyncio.sleep(random.uniform(1.5, 3.5))

            except Exception as e:
                print(f"⚠️ Fatal Error: {e}")
                break

        print(f"✅ Completed '{keyword}' — {len(job_postings)} jobs found.\n")
        return job_postings


# =============================================
# EXPORTER (Excel + Word + PDF)
# =============================================
class Exporter:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.hyperlink_style = ParagraphStyle(
            'Hyperlink',
            parent=self.styles['Normal'],
            fontSize=11,
            textColor=blue,
            underline=True,
            leading=14
        )
        self.normal_style = ParagraphStyle("Normal", fontSize=11, leading=14)

    @staticmethod
    def remove_empty_sheet(wb):
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if sheet_name == "Sheet" and all([cell.value is None for row in ws.iter_rows() for cell in row]):
                wb.remove(ws)

    def save_excel(self, file_path, sheet_name, jobs):
        folder = os.path.dirname(file_path)
        if folder:
            os.makedirs(folder, exist_ok=True)

        if os.path.exists(file_path):
            wb = load_workbook(file_path)
            if sheet_name in wb.sheetnames:
                del wb[sheet_name]
        else:
            wb = Workbook()

        ws = wb.create_sheet(title=sheet_name)
        headers = ["Title", "Company", "Location", "Date Posted", "Keyword", "Job Link", "Mobile Link"]
        ws.append(headers)

        for job in jobs:
            row = [
                job["Title"],
                job["Company"],
                job["Location"],
                job["Date Posted"],
                job["Keyword"],
                job["Job Link"],
                job["Mobile Link"]
            ]
            ws.append(row)
            ws.cell(row=ws.max_row, column=6).hyperlink = job["Job Link"]
            ws.cell(row=ws.max_row, column=6).font = Font(color="0000FF", underline="single")
            ws.cell(row=ws.max_row, column=7).hyperlink = job["Mobile Link"]
            ws.cell(row=ws.max_row, column=7).font = Font(color="0000FF", underline="single")

        self.remove_empty_sheet(wb)
        wb.save(file_path)

    def save_keyword_files(self, keyword, jobs):
        date_code = datetime.now().strftime("%Y%m%d")
        folder = keyword.replace(" ", "")
        os.makedirs(folder, exist_ok=True)

        # Per-keyword Excel
        kw_excel = os.path.join(folder, f"{folder}_{date_code}_jobs.xlsx")
        self.save_excel(kw_excel, sheet_name=keyword.replace(" ", ""), jobs=jobs)

        # Per-keyword Word + PDF
        kw_docx = os.path.join(folder, f"{folder}_{date_code}_jobs.docx")
        kw_pdf = os.path.join(folder, f"{folder}_{date_code}_jobs.pdf")
        self.save_keyword_documents(jobs, kw_docx, kw_pdf)

        return kw_excel, kw_docx, kw_pdf

    def save_keyword_documents(self, jobs, docx_file, pdf_file):
        doc = Document()

        for job in jobs:
            p = doc.add_paragraph()
            p.add_run("🧾 Title: ").bold = True
            p.add_run(job["Title"])

            p = doc.add_paragraph()
            p.add_run("🏢 Company: ").bold = True
            p.add_run(job["Company"])

            p = doc.add_paragraph()
            p.add_run("📍 Location: ").bold = True
            p.add_run(job["Location"])

            p = doc.add_paragraph()
            p.add_run("📅 Posted: ").bold = True
            p.add_run(job["Date Posted"])

            p = doc.add_paragraph()
            p.add_run("🔗 Link: ").bold = True
            self.add_hyperlink(p, job["Mobile Link"], job["Mobile Link"])

            p = doc.add_paragraph()
            p.add_run("🔗 Web Link: ").bold = True
            self.add_hyperlink(p, job["Job Link"], job["Job Link"])

            doc.add_paragraph("-" * 80)

        doc.save(docx_file)
        self.convert_docx_to_pdf(docx_file, pdf_file)

    @staticmethod
    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    @staticmethod
    def convert_docx_to_pdf(docx_file, pdf_file):
        doc = Document(docx_file)
        story = []
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        hyperlink_style = ParagraphStyle('Hyperlink', parent=normal_style, fontSize=11, textColor=blue, underline=True)
        for para in doc.paragraphs:
            text = para.text
            style_to_use = hyperlink_style if "🔗" in text else normal_style
            story.append(Paragraph(text, style_to_use))
            story.append(Spacer(1, 12))
        pdf = SimpleDocTemplate(pdf_file, pagesize=LETTER)
        pdf.build(story)


# =============================================
# SCRAPER RUNNER
# =============================================
class ScraperRunner:
    def __init__(self, db: DBClient, scraper: LinkedInScraper, exporter: Exporter, keywords: list):
        self.db = db
        self.scraper = scraper
        self.exporter = exporter
        self.keywords = keywords

    async def run(self, client):
        for keyword in self.keywords:
            run_id = self.db.log_run_start(keyword, SOURCE_PORTAL)
            self.db.start_transaction()

            try:
                jobs = await self.scraper.fetch_jobs_for_keyword(client, keyword)
                if not jobs:
                    print(f"ℹ️ No jobs found for keyword '{keyword}'")
                    self.db.log_run_end(run_id, 0)
                    self.db.rollback_transaction()
                    continue

                # Export Excel + Word + PDF
                date_code = datetime.now().strftime("%Y%m%d")
                main_excel = f"Job_Extract_{date_code}.xlsx"
                self.exporter.save_excel(main_excel, keyword.replace(" ", ""), jobs)
                kw_excel, kw_docx, kw_pdf = self.exporter.save_keyword_files(keyword, jobs)

                # Insert into DB
                for job in jobs:
                    self.db.upsert_master(job, SOURCE_PORTAL)

                # Commit all changes
                self.db.commit_transaction()
                self.db.log_run_end(run_id, len(jobs))
                print(f"🎉 Completed keyword '{keyword}' — {len(jobs)} jobs inserted.\n")

            except Exception as e:
                print(f"❌ Transaction rolled back for keyword '{keyword}' due to error: {e}")
                self.db.rollback_transaction()


# =============================================
# MAIN
# =============================================
async def main():
    db = DBClient()
    db.archive_master_to_history()
    db.clear_master()

    scraper = LinkedInScraper()
    exporter = Exporter()
    runner = ScraperRunner(db, scraper, exporter, KEYWORDS)

    async with httpx.AsyncClient(timeout=45.0) as client:
        await runner.run(client)

    db.close()


# ---------------------------------------------
# RUN
# ---------------------------------------------
if __name__ == "__main__":
    asyncio.run(main())
