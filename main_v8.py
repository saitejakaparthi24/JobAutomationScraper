# =============================================
#   BEAST MODE LINKEDIN SCRAPER (Daily Excel + Hyperlinks)
#   Console Output + Excel + Last 7 Days + Sorted + Folder Excel
# =============================================

import asyncio
import httpx
from bs4 import BeautifulSoup
import re
import random
import os
from datetime import datetime, timedelta
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
from db_client import DBClient
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import LETTER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.colors import blue

# ---------------------------------------------
# CONFIG
# ---------------------------------------------
KEYWORDS = ["Backend Developer", "SAP SD Consultant"]
LOCATION = "United States"
MAX_RETRY = 3
SOURCE_PORTAL = "LinkedIn"

# ---------------------------------------------
# DATE FILTER (Last 7 days)
# ---------------------------------------------
def posted_within_last_week(date_str):
    try:
        post_date = datetime.fromisoformat(date_str.replace("Z", ""))
        return post_date >= datetime.now() - timedelta(days=7)
    except:
        return False

# ---------------------------------------------
# CLEAN JOB ID → Mobile Link
# ---------------------------------------------
def get_mobile_link(job_link):
    match = re.search(r"/jobs/view/.*?-(\d+)", job_link)
    return f"https://www.linkedin.com/jobs/view/{match.group(1)}" if match else job_link

# ---------------------------------------------
# LOAD PREVIOUS DAY JOB IDS FOR DEDUPLICATION
# ---------------------------------------------
def load_previous_ids(keyword):
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y%m%d")
    filename = f"{keyword.replace(' ', '')}/{keyword.replace(' ', '')}_{yesterday}.xlsx"
    if not os.path.exists(filename):
        return set()

    previous_ids = set()
    try:
        wb = load_workbook(filename)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            job_link = row[5]  # Job Link column
            m = re.search(r"/jobs/view/.*?-(\d+)", job_link)
            if m:
                previous_ids.add(m.group(1))
    except:
        pass
    return previous_ids

# ---------------------------------------------
# SCRAPE PER KEYWORD
# ---------------------------------------------
async def fetch_jobs_for_keyword(client, keyword):
    url = "https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                      "(KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36"
    }

    job_postings = []
    seen_ids_today = set()
    seen_ids_yesterday = load_previous_ids(keyword)

    print(f"\n🚀 Starting scrape for keyword: {keyword}")
    print(f"📌 Loaded {len(seen_ids_yesterday)} previous job IDs (for dedupe)")

    page = 0
    retry = 0

    while True:
        params = {
            "keywords": keyword,
            "location": LOCATION,
            "sortBy": "R",          # Most recent
            "f_TPR": "r86400",     # Last 7 days
            "start": page * 25
        }

        print(f"🔄 Fetching page {page}...")

        try:
            resp = await client.get(url, headers=headers, params=params)

            if resp.status_code == 429:
                wait_time = random.randint(45, 90)
                print(f"🚫 Rate limited (429). Sleeping {wait_time} sec...")
                await asyncio.sleep(wait_time)
                continue

            if resp.status_code != 200:
                print(f"❌ HTTP {resp.status_code} on page {page}")
                retry += 1
                if retry > MAX_RETRY:
                    print("❌ Too many failures. Stopping scraper.")
                    break
                await asyncio.sleep(5)
                continue

            retry = 0
            soup = BeautifulSoup(resp.content, "lxml")
            job_cards = soup.select("li")

            if not job_cards:
                print(f"ℹ️ No more jobs found for: {keyword}")
                break

            page += 1

            for job in job_cards:
                try:
                    title_tag = job.find("h3")
                    company_tag = job.find("h4")
                    location_tag = job.find("span", class_="job-search-card__location")
                    time_tag = job.find("time")
                    link_tag = job.find("a", href=True)

                    if not all([title_tag, company_tag, location_tag, time_tag, link_tag]):
                        continue

                    job_link = link_tag["href"]
                    mobile_link = get_mobile_link(job_link)

                    job_id_match = re.search(r"/jobs/view/.*?-(\d+)", job_link)
                    if not job_id_match:
                        continue
                    job_id = job_id_match.group(1)

                    if job_id in seen_ids_today or job_id in seen_ids_yesterday:
                        continue

                    seen_ids_today.add(job_id)

                    title = title_tag.text.strip()
                    company = company_tag.text.strip()
                    location = location_tag.text.strip()
                    date_posted = time_tag.get("datetime", "")

                    if not posted_within_last_week(date_posted):
                        continue

                    job_postings.append({
                        "Title": title,
                        "Company": company,
                        "Location": location,
                        "Date Posted": date_posted,
                        "Keyword": keyword,
                        "Job Link": job_link,
                        "Mobile Link": mobile_link
                    })

                except Exception as e:
                    print(f"⚠️ Parse Error: {e}")
                    continue

            await asyncio.sleep(random.uniform(1.5, 3.5))

        except Exception as e:
            print(f"⚠️ Fatal Error: {e}")
            break

    print(f"✅ Completed '{keyword}' — {len(job_postings)} jobs found.\n")
    return job_postings

# ---------------------------------------------
# SAVE EXCEL WITH HYPERLINKS
# ---------------------------------------------
def save_excel(file_path, sheet_name, jobs):
    folder = os.path.dirname(file_path)
    if folder:
        os.makedirs(folder, exist_ok=True)

    if os.path.exists(file_path):
        wb = load_workbook(file_path)
        if sheet_name in wb.sheetnames:
            del wb[sheet_name]
    else:
        wb = Workbook()

    ws = wb.create_sheet(title=sheet_name)

    headers = ["Title", "Company", "Location", "Date Posted", "Keyword", "Job Link", "Mobile Link"]
    ws.append(headers)

    for job in jobs:
        row = [
            job["Title"],
            job["Company"],
            job["Location"],
            job["Date Posted"],
            job["Keyword"],
            job["Job Link"],
            job["Mobile Link"]
        ]
        ws.append(row)

        # Add hyperlink for Job Link column (6th column)
        job_link_cell = ws.cell(row=ws.max_row, column=6, value=job["Job Link"])
        job_link_cell.hyperlink = job["Job Link"]
        job_link_cell.font = Font(color="0000FF", underline="single")

        # Add hyperlink for Mobile Link column (7th column)
        mobile_link_cell = ws.cell(row=ws.max_row, column=7, value=job["Mobile Link"])
        mobile_link_cell.hyperlink = job["Mobile Link"]
        mobile_link_cell.font = Font(color="0000FF", underline="single")

    # Remove dummy default sheet if truly empty
    remove_empty_sheet(wb)

    wb.save(file_path)

def remove_empty_sheet(wb):
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # Consider sheet empty if it has no data in any cell
        if sheet_name == "Sheet":
            if all([cell.value is None for row in ws.iter_rows() for cell in row]):
                wb.remove(ws)

def add_hyperlink(paragraph, text, url):
    """
    Create a clickable hyperlink in a python-docx paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # blue
    rPr.append(color)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text

    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def append_to_keyword_documents(keyword, job):
    from docx import Document

    date_code = datetime.now().strftime("%Y%m%d")
    safe_keyword = keyword.replace(" ", "")
    docx_file = f"{safe_keyword}_{date_code}.docx"
    pdf_file  = f"{safe_keyword}_{date_code}.pdf"

    # Load or create doc
    if os.path.exists(docx_file):
        doc = Document(docx_file)
    else:
        doc = Document()

    # Title
    p = doc.add_paragraph()
    p.add_run("🧾 Title: ").bold = True
    p.add_run(job["Title"])

    p = doc.add_paragraph()
    p.add_run("🏢 Company: ").bold = True
    p.add_run(job["Company"])

    p = doc.add_paragraph()
    p.add_run("📍 Location: ").bold = True
    p.add_run(job["Location"])

    p = doc.add_paragraph()
    p.add_run("📅 Posted: ").bold = True
    p.add_run(job["Date Posted"])

    p = doc.add_paragraph()
    p.add_run("🔗 Link: ").bold = True
    add_hyperlink(p, job["Mobile Link"], job["Mobile Link"])

    p = doc.add_paragraph()
    p.add_run("🔗 Web Link: ").bold = True
    add_hyperlink(p, job["Job Link"], job["Job Link"])

    doc.add_paragraph("-" * 80)
    doc.save(docx_file)

    # ----- PDF FROM DOCX -----
    convert_docx_to_pdf(docx_file, pdf_file)

def convert_docx_to_pdf(docx_file, pdf_file):
    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.colors import blue

    doc = Document(docx_file)

    # -----------------------------
    # ✅ FIX 3 — Add hyperlink style
    # -----------------------------
    styles = getSampleStyleSheet()
    hyperlink_style = ParagraphStyle(
        'Hyperlink',
        parent=styles['Normal'],
        fontSize=11,
        textColor=blue,
        underline=True,
        leading=14
    )
    normal_style = ParagraphStyle("Normal", fontSize=11, leading=14)
    # -----------------------------

    pdf = SimpleDocTemplate(pdf_file, pagesize=LETTER)
    story = []

    for para in doc.paragraphs:
        text = para.text

        # Detect hyperlinks
        if text.startswith("🔗 Web Link: "):
            url = text.replace("🔗 Web Link: ", "").strip()
            text = f'🔗 Web Link: <a href="{url}">{url}</a>'
            style_to_use = hyperlink_style

        elif text.startswith("🔗 Link: "):
            url = text.replace("🔗 Link: ", "").strip()
            text = f'🔗 Link: <a href="{url}">{url}</a>'
            style_to_use = hyperlink_style

        else:
            style_to_use = normal_style

        story.append(Paragraph(text, style_to_use))
        story.append(Spacer(1, 12))

    pdf.build(story)



# ---------------------------------------------
# MAIN
# ---------------------------------------------
async def main():
    db = DBClient()
    date_code = datetime.now().strftime("%Y%m%d")
    main_excel = f"Job_Extract_{date_code}.xlsx"

    # 1️⃣ Archive yesterday's data
    print("📦 Archiving yesterday's job_master data to job_daily_history...")
    db.archive_master_to_history()

    # 2️⃣ Clear master for today
    print("🧹 Clearing job_master table for today's run...")
    db.clear_master()

    async with httpx.AsyncClient(timeout=45.0) as client:

        for keyword in KEYWORDS:
            run_id = db.log_run_start(keyword, SOURCE_PORTAL)
            jobs = await fetch_jobs_for_keyword(client, keyword)

            save_excel(main_excel, keyword, jobs)
            folder_file = f"{keyword.replace(' ', '')}/LinkedIn_{keyword.replace(' ', '')}_Jobs_{date_code}.xlsx"
            save_excel(folder_file, keyword, jobs)

            print("\n========================")
            print(f"📢 RESULTS FOR {keyword}")
            print("========================\n")
            for job in jobs:
                print("🧾 Title   :", job["Title"])
                print("🏢 Company :", job["Company"])
                print("📍 Location:", job["Location"])
                print("📅 Posted  :", job["Date Posted"])
                print("🔑 Keyword :", job["Keyword"])
                print("🔗 Link    :", job["Mobile Link"])
                print("-" * 80)
                doc_text = (
                    f"🧾 Title   : {job['Title']}\n"
                    f"🏢 Company : {job['Company']}\n"
                    f"📍 Location: {job['Location']}\n"
                    f"📅 Posted  : {job['Date Posted']}\n"
                    f"🔑 Keyword : {job['Keyword']}\n"
                    f"🔗 Link    : {job['Mobile Link']}\n"
                    f"{'-' * 80}\n"
                )

                append_to_keyword_documents(job["Keyword"], job)

            print(f"\n📁 Saved to main Excel: {main_excel}")
            print(f"📁 Saved to folder Excel: {folder_file}")
            print("🎉 Scraping Completed Successfully!\n")


            # ---- DATABASE INSERT ----
            for job in jobs:
                job_id = db.upsert_master(job, SOURCE_PORTAL)

            db.log_run_end(run_id, len(jobs))
            print(f"🎉 {len(jobs)} Jobs Inserted and Run Completed Successfully for {keyword}\n")
            db.cleanup_history()
    db.close()

# ---------------------------------------------
# RUN
# ---------------------------------------------
if __name__ == "__main__":
    asyncio.run(main())
